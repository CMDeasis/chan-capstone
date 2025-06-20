"""
Test script for the DPA Compliance System
Creates sample documents and tests the compliance checker
"""

from faker import Faker
import os
from docx import Document
from dpa_compliance_checker import DPAComplianceChecker
from report_generator import DPAReportGenerator

# Initialize Faker for Philippine locale
fake = Faker(['en_PH', 'en_US'])

def create_sample_documents():
    """Create sample documents with PII for testing"""
    
    # Ensure data/input directory exists
    os.makedirs('data/input', exist_ok=True)
    
    # Sample 1: Document with violations (no consent, has sensitive data)
    sample1_text = f"""
    EMPLOYEE RECORD
    
    Name: {fake.name()}
    TIN: {fake.random_number(digits=3)}-{fake.random_number(digits=3)}-{fake.random_number(digits=3)}-{fake.random_number(digits=3)}
    SSS Number: {fake.random_number(digits=2)}-{fake.random_number(digits=7)}-{fake.random_number(digits=1)}
    Phone: +63{fake.random_number(digits=10)}
    Email: {fake.email()}
    
    Medical Information:
    - Has diabetes
    - Underwent surgery in 2023
    - Takes medication for hypertension
    
    Religious Affiliation: Catholic
    Political Party: Liberal Party
    
    Salary: PHP {fake.random_number(digits=6)}
    Bank Account: {fake.iban()}
    
    This information is collected for employment purposes.
    """
    
    # Save as DOCX
    doc1 = Document()
    doc1.add_paragraph(sample1_text)
    doc1.save('data/input/sample_violations.docx')
    
    # Sample 2: Compliant document (has consent and purpose)
    sample2_text = f"""
    CUSTOMER REGISTRATION FORM
    
    CONSENT STATEMENT:
    I, {fake.name()}, hereby give my consent for the processing of my personal information 
    for the purpose of account registration and customer service.
    
    Personal Information:
    Name: {fake.name()}
    Phone: +63{fake.random_number(digits=10)}
    Email: {fake.email()}
    Address: {fake.address()}
    
    PURPOSE: This information will be used solely for account management and 
    customer communication purposes in accordance with the Data Privacy Act of 2012.
    
    Data Subject Signature: ________________
    Date: {fake.date()}
    """
    
    # Save as DOCX
    doc2 = Document()
    doc2.add_paragraph(sample2_text)
    doc2.save('data/input/sample_compliant.docx')
    
    # Sample 3: Mixed compliance (some consent, some violations)
    sample3_text = f"""
    HEALTHCARE PATIENT RECORD
    
    Patient Name: {fake.name()}
    TIN: {fake.random_number(digits=3)}-{fake.random_number(digits=3)}-{fake.random_number(digits=3)}-{fake.random_number(digits=3)}
    PhilHealth: {fake.random_number(digits=2)}-{fake.random_number(digits=9)}-{fake.random_number(digits=1)}
    
    I consent to the processing of my medical information for treatment purposes.
    
    Medical History:
    - Diagnosed with cancer in 2022
    - Currently undergoing chemotherapy
    - Mental health: Depression, anxiety
    - Previous surgeries: Appendectomy (2020)
    
    Emergency Contact:
    Name: {fake.name()}
    Phone: +63{fake.random_number(digits=10)}
    
    Insurance Information:
    Provider: {fake.company()}
    Policy Number: {fake.random_number(digits=10)}
    
    Note: This medical record contains sensitive health information.
    """
    
    # Save as DOCX
    doc3 = Document()
    doc3.add_paragraph(sample3_text)
    doc3.save('data/input/sample_mixed.docx')
    
    print("Sample documents created:")
    print("1. data/input/sample_violations.docx - Document with multiple violations")
    print("2. data/input/sample_compliant.docx - Compliant document")
    print("3. data/input/sample_mixed.docx - Mixed compliance document")

def test_compliance_checker():
    """Test the compliance checker with sample documents"""
    
    checker = DPAComplianceChecker()
    report_generator = DPAReportGenerator()
    
    # Test documents
    test_files = [
        ('data/input/sample_violations.docx', 'High Violations Document'),
        ('data/input/sample_compliant.docx', 'Compliant Document'),
        ('data/input/sample_mixed.docx', 'Mixed Compliance Document')
    ]
    
    for file_path, description in test_files:
        if os.path.exists(file_path):
            print(f"\n{'='*60}")
            print(f"Testing: {description}")
            print(f"File: {file_path}")
            print('='*60)
            
            # Read document content
            doc = Document(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
            
            # Analyze compliance
            report = checker.analyze_document(text, description)
            
            # Generate summary
            summary = checker.generate_summary(report)
            
            # Print results
            print(f"Compliance Status: {summary['status']}")
            print(f"Risk Level: {summary['risk_level']}")
            print(f"Total Violations: {summary['total_violations']}")
            print(f"PII Found: {summary['pii_found']}")
            print(f"Sensitive PII: {summary['sensitive_pii_found']}")
            
            if summary['key_issues']:
                print("\nKey Issues:")
                for issue in summary['key_issues']:
                    print(f"  - {issue}")
            
            if summary['top_recommendations']:
                print("\nTop Recommendations:")
                for rec in summary['top_recommendations']:
                    print(f"  - {rec}")
            
            # Generate PDF report
            report_filename = f"test_report_{description.lower().replace(' ', '_')}.pdf"
            report_path = os.path.join('data', 'output', report_filename)
            os.makedirs(os.path.dirname(report_path), exist_ok=True)
            
            report_generator.generate_report(report, report_path)
            print(f"\nDetailed PDF report saved: {report_path}")
        else:
            print(f"File not found: {file_path}")

def test_pii_detection():
    """Test PII detection with various text samples"""
    
    from pii_detector import EnhancedPIIDetector
    
    detector = EnhancedPIIDetector()
    
    test_texts = [
        "Juan Dela Cruz, TIN: 123-456-789-000, has diabetes and is Catholic.",
        "Contact Maria Santos at +639171234567 or maria@email.com",
        "SSS Number: 12-3456789-0, PhilHealth: 12-345678901-2",
        "Patient diagnosed with cancer, takes medication for hypertension",
        "Salary: PHP 50000, Bank Account: PH12345678901234567890"
    ]
    
    print("\n" + "="*60)
    print("PII DETECTION TEST")
    print("="*60)
    
    for i, text in enumerate(test_texts, 1):
        print(f"\nTest {i}: {text}")
        print("-" * 40)
        
        pii_results = detector.detect_pii(text)
        categorized = detector.categorize_pii(pii_results)
        
        print(f"Total PII: {categorized['total_pii_count']}")
        print(f"Sensitive PII: {categorized['sensitive_count']}")
        
        if pii_results:
            for pii in pii_results:
                print(f"  - {pii['entity_type']}: {pii['text']} (Confidence: {pii['confidence']:.2f})")
        else:
            print("  No PII detected")

if __name__ == "__main__":
    print("DPA Compliance System Test")
    print("=" * 50)
    
    # Create sample documents
    print("\n1. Creating sample documents...")
    create_sample_documents()
    
    # Test PII detection
    print("\n2. Testing PII detection...")
    test_pii_detection()
    
    # Test compliance checker
    print("\n3. Testing compliance checker...")
    test_compliance_checker()
    
    print("\n" + "="*60)
    print("TEST COMPLETED")
    print("="*60)
    print("Check the data/output directory for generated PDF reports.")
    print("You can now run the Flask app with: python app.py")
